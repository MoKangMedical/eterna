from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "eterna-business-plan-mckinsey-style.md"
POSTER = ROOT / "frontend" / "assets" / "eterna-demo-v2-poster.jpg"
OUTPUT_DIR = ROOT / "output" / "doc"
OUT_DOCX = OUTPUT_DIR / "eterna-business-plan-external.docx"
OUT_SCRIPT = OUTPUT_DIR / "build_external_bp_doc.py"


COLORS = {
    "navy": "0D1B2A",
    "navy_2": "1B2A41",
    "line": "D7DEE7",
    "muted": "5B6673",
    "accent": "D98B67",
    "teal": "5C9FA5",
    "white": "FFFFFF",
    "light": "F6F8FB",
}

DOC_FONT = "Arial Unicode MS"


def set_run_font(run, east_asia: str = DOC_FONT, latin: str = DOC_FONT) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_style_font(style, east_asia: str = DOC_FONT, latin: str = DOC_FONT, size: int | None = None) -> None:
    style.font.name = latin
    if size:
        style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    styles = doc.styles
    set_style_font(styles["Normal"], east_asia=DOC_FONT, latin=DOC_FONT, size=10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(COLORS["navy_2"])
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Title"], east_asia=DOC_FONT, latin=DOC_FONT, size=26)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(COLORS["navy"])
    styles["Title"].paragraph_format.space_after = Pt(8)

    for name, size, color in [("Heading 1", 17, COLORS["navy"]), ("Heading 2", 13, COLORS["accent"]), ("Heading 3", 11, COLORS["teal"])]:
      set_style_font(styles[name], east_asia=DOC_FONT, latin=DOC_FONT, size=size)
      styles[name].font.bold = True
      styles[name].font.color.rgb = RGBColor.from_string(color)
      styles[name].paragraph_format.space_before = Pt(12)
      styles[name].paragraph_format.space_after = Pt(6)

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)
    return doc


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(0)
    run = p.add_run("念念 Eterna")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("商业计划书")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["accent"])
    set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("对外正式版 | McKinsey-style")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(18)
    p.space_after = Pt(18)
    run = p.add_run("Prepared for investors and strategic partners")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    set_run_font(run, east_asia=DOC_FONT, latin=DOC_FONT)

    if POSTER.exists():
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(POSTER), width=Mm(128))

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.space_before = Pt(16)
    lines = [
        "项目名称：念念 Eterna",
        "文档版本：v1.1",
        "日期：2026-04-09",
        "文档属性：商业计划书 / 对外沟通稿",
    ]
    for idx, line in enumerate(lines):
        run = info.add_run(line)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(COLORS["navy_2"])
        set_run_font(run)
        if idx < len(lines) - 1:
            run.add_break()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.space_before = Pt(18)
    note_format = note.paragraph_format
    note_format.left_indent = Mm(20)
    note_format.right_indent = Mm(20)
    run = note.add_run("机密提示：本文件用于潜在投资人与战略合作方沟通，不构成历史经营数据承诺。文中财务与增长数据如无特别说明，均为经营假设。")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    set_run_font(run, east_asia=DOC_FONT, latin=DOC_FONT)

    doc.add_page_break()


def extract_section_titles(lines: list[str]) -> list[str]:
    titles: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            titles.append(stripped[3:].strip())
    return titles


def add_toc_page(doc: Document, titles: list[str]) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("目录")
    set_run_font(run)

    for title in titles:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Mm(6)
        run = para.add_run(title)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(COLORS["navy_2"])
        set_run_font(run, east_asia=DOC_FONT, latin=DOC_FONT)

    note = doc.add_paragraph()
    run = note.add_run("注：本版为对外排版稿，后续如需送印或正式提交，可再做目录页码与法务页优化。")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    set_run_font(run)
    doc.add_page_break()


def parse_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    return rows, i


def add_md_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    set_run_font(run, east_asia=DOC_FONT, latin=DOC_FONT)
                    if r_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(COLORS["white"])
            if r_idx == 0:
                set_cell_shading(cell, COLORS["navy"])


def add_bullet_list(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        run.font.size = Pt(10.2)
        set_run_font(run, east_asia=DOC_FONT, latin=DOC_FONT)


def add_paragraph_text(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Normal")
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    set_run_font(run, east_asia=DOC_FONT, latin=DOC_FONT)


def render_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    start = 0
    for idx, line in enumerate(lines):
        if line.strip().startswith("## 0."):
            start = idx
            break
    lines = lines[start:]

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("# "):
            para = doc.add_paragraph(style="Heading 1")
            run = para.add_run(stripped[2:].strip())
            set_run_font(run)
            i += 1
            continue
        if stripped.startswith("## "):
            para = doc.add_paragraph(style="Heading 1")
            run = para.add_run(stripped[3:].strip())
            set_run_font(run)
            i += 1
            continue
        if stripped.startswith("### "):
            para = doc.add_paragraph(style="Heading 2")
            run = para.add_run(stripped[4:].strip())
            set_run_font(run)
            i += 1
            continue
        if stripped.startswith("#### "):
            para = doc.add_paragraph(style="Heading 3")
            run = para.add_run(stripped[5:].strip())
            set_run_font(run)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, next_i = parse_table(lines, i)
            add_md_table(doc, rows)
            i = next_i
            continue
        if re.match(r"^- ", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^- ", lines[i].strip()):
                items.append(re.sub(r"^- ", "", lines[i].strip()))
                i += 1
            add_bullet_list(doc, items, numbered=False)
            continue
        if re.match(r"^\d+\. ", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^\d+\. ", lines[i].strip()):
                items.append(re.sub(r"^\d+\. ", "", lines[i].strip()))
                i += 1
            add_bullet_list(doc, items, numbered=True)
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|") or re.match(r"^- ", nxt) or re.match(r"^\d+\. ", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        add_paragraph_text(doc, " ".join(paragraph_lines))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    raw_lines = md_text.splitlines()
    titles = extract_section_titles(raw_lines)

    doc = build_document()
    add_cover(doc)
    add_toc_page(doc, titles)
    render_markdown(doc, md_text)
    doc.save(OUT_DOCX)
    shutil.copyfile(__file__, OUT_SCRIPT)
    print(f"Wrote {OUT_DOCX}")
    print(f"Copied source to {OUT_SCRIPT}")


if __name__ == "__main__":
    main()
